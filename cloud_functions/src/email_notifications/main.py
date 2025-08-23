import asyncio
import base64
import json
import os
from io import BytesIO
from pathlib import Path
from typing import Any, NotRequired, TypedDict

import functions_framework
import httpx
from cloudevents.http import CloudEvent
from docx import Document
from jinja2 import Environment, FileSystemLoader
from markdown import markdown
from packages.db.src.tables import GrantApplication, Organization, OrganizationUser, Project
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
from sqlalchemy.orm import sessionmaker

logger = __import__("logging").getLogger(__name__)


class EmailResponse(TypedDict):
    status: str
    message: str


class ApplicationData(TypedDict):
    id: str
    title: str
    text: str
    created_at: str


class UserData(TypedDict):
    email: str
    name: str


class ProjectData(TypedDict):
    id: str
    name: str


class ApplicationDataResponse(TypedDict):
    application: ApplicationData
    user: UserData
    project: ProjectData


class GrantAlertTemplateData(TypedDict):
    grants: list[dict[str, Any]]
    frequency: str
    unsubscribe_url: str
    subscription_id: NotRequired[str]


class GrantAlertNotification(TypedDict):
    email: str
    template_data: GrantAlertTemplateData
    template_type: NotRequired[str]


class SubscriptionVerificationData(TypedDict):
    email: str
    subscription_id: str
    verification_token: str
    search_params: NotRequired[dict[str, Any]]
    frequency: NotRequired[str]
    template_type: NotRequired[str]


class ApplicationNotification(TypedDict):
    application_id: str
    template_type: NotRequired[str]


class ResendAttachment(TypedDict):
    filename: str
    content: str


template_dir = Path(__file__).parent / "templates"
jinja_env = Environment(loader=FileSystemLoader(template_dir), autoescape=True)


async def get_application_data(application_id: str) -> ApplicationDataResponse:
    db_connection_string = os.environ.get("DATABASE_CONNECTION_STRING")
    if not db_connection_string:
        raise ValueError("DATABASE_CONNECTION_STRING not configured")

    engine = create_async_engine(db_connection_string)
    async_session_maker = sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)  # type: ignore[call-overload]

    async with async_session_maker() as session:
        app_result = await session.execute(
            select(GrantApplication, Project)
            .join(Project, GrantApplication.project_id == Project.id)
            .where(GrantApplication.id == application_id)
        )
        app_row = app_result.first()

        if not app_row:
            raise ValueError(f"Application {application_id} not found")

        application, project = app_row

        org_result = await session.execute(
            select(OrganizationUser)
            .where(OrganizationUser.organization_id == project.organization_id)
            .where(OrganizationUser.role == "OWNER")
            .limit(1)
        )
        org_user = org_result.scalar_one_or_none()

        org_result = await session.execute(select(Organization).where(Organization.id == project.organization_id))
        organization = org_result.scalar_one()

        return ApplicationDataResponse(
            application=ApplicationData(
                id=str(application.id),
                title=application.title,
                text=application.text,
                created_at=application.created_at.isoformat(),
            ),
            user=UserData(
                email=organization.contact_email or (org_user.firebase_uid if org_user else "user@example.com"),
                name=organization.contact_person_name or organization.name,
            ),
            project=ProjectData(
                id=str(project.id),
                name=project.name,
            ),
        )


def markdown_to_docx(markdown_text: str) -> bytes:
    markdown(markdown_text, extensions=["tables", "fenced_code", "nl2br"])

    doc = Document()
    doc.add_heading("Grant Application", 0)

    lines = markdown_text.split("\n")
    current_paragraph = None

    for line in lines:
        stripped_line = line.strip()

        if stripped_line.startswith("# "):
            doc.add_heading(stripped_line[2:], level=1)
            current_paragraph = None
        elif stripped_line.startswith("## "):
            doc.add_heading(stripped_line[3:], level=2)
            current_paragraph = None
        elif stripped_line.startswith("### "):
            doc.add_heading(stripped_line[4:], level=3)
            current_paragraph = None
        elif stripped_line.startswith(("- ", "* ")):
            doc.add_paragraph(stripped_line[2:], style="List Bullet")
            current_paragraph = None
        elif stripped_line.startswith("1. ") or (
            len(stripped_line) > 2 and stripped_line[0].isdigit() and stripped_line[1:3] == ". "
        ):
            doc.add_paragraph(stripped_line[stripped_line.index(". ") + 2 :], style="List Number")
            current_paragraph = None
        elif stripped_line:
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph(stripped_line)
            else:
                current_paragraph.add_run(" " + stripped_line)
        else:
            current_paragraph = None

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


async def send_resend_email(
    to_email: str, subject: str, html: str, attachments: list[ResendAttachment]
) -> dict[str, Any]:
    resend_api_key = os.environ.get("RESEND_API_KEY")
    if not resend_api_key:
        raise ValueError("RESEND_API_KEY not configured")

    async with httpx.AsyncClient() as client:
        response = await client.post(
            "https://api.resend.com/emails",
            headers={
                "Authorization": f"Bearer {resend_api_key}",
                "Content-Type": "application/json",
            },
            json={
                "from": "GrantFlow <notifications@grantflow.ai>",
                "to": [to_email],
                "subject": subject,
                "html": html,
                "attachments": attachments,
            },
            timeout=30.0,
        )

        if response.status_code != 200:
            raise Exception(f"Resend API error: {response.status_code} - {response.text}")

        return response.json()  # type: ignore[no-any-return]


async def send_subscription_verification_email(notification_data: SubscriptionVerificationData) -> EmailResponse:
    try:
        email = notification_data["email"]
        subscription_id = notification_data["subscription_id"]
        verification_token = notification_data["verification_token"]
        search_params = notification_data.get("search_params", {})
        frequency = notification_data.get("frequency", "daily")

        site_url = os.environ.get("SITE_URL", "https://grantflow.ai")
        verification_url = f"{site_url}/public/grants/verify/{verification_token}"

        template = jinja_env.get_template("subscription_verification.html")
        html_content = template.render(
            verification_url=verification_url,
            frequency=frequency,
            search_query=search_params.get("query"),
            category=search_params.get("category"),
            min_amount=search_params.get("min_amount"),
            max_amount=search_params.get("max_amount"),
        )

        subject = "Verify Your GrantFlow Grant Alert Subscription"

        await send_resend_email(email, subject, html_content, [])

        logger.info(
            "Subscription verification email sent",
            email=email,
            subscription_id=subscription_id,
        )

        return EmailResponse(status="success", message="Verification email sent successfully")

    except Exception as e:
        logger.exception("Failed to send subscription verification email")
        return EmailResponse(status="error", message=f"Failed to send verification email: {e!s}")


async def send_grant_alert_email(notification_data: GrantAlertNotification) -> EmailResponse:
    try:
        email = notification_data["email"]
        template_data = notification_data["template_data"]

        template = jinja_env.get_template("grant_alert.html")
        html_content = template.render(
            grants=template_data["grants"],
            frequency=template_data["frequency"],
            unsubscribe_url=template_data["unsubscribe_url"],
        )

        grant_count = len(template_data["grants"])
        subject = f"🎯 {grant_count} New Grant{'s' if grant_count != 1 else ''} Available"

        await send_resend_email(email, subject, html_content, [])

        logger.info(
            "Grant alert email sent",
            email=email,
            grant_count=grant_count,
        )

        return EmailResponse(status="success", message="Grant alert sent successfully")

    except Exception as e:
        logger.exception("Failed to send grant alert email")
        return EmailResponse(status="error", message=f"Failed to send grant alert: {e!s}")


def _extract_pubsub_data(cloud_event: CloudEvent) -> dict[str, Any]:
    event_data = cloud_event.data
    if event_data is None:
        attributes = cloud_event.get_attributes()
        if "data" in attributes:
            event_data = attributes["data"]

    if isinstance(event_data, dict) and "message" in event_data:
        pubsub_message = event_data["message"]
        if "data" in pubsub_message:
            message_data = base64.b64decode(pubsub_message["data"]).decode("utf-8")
            return json.loads(message_data)  # type: ignore[no-any-return]
        raise ValueError("No data in Pub/Sub message")

    raise ValueError("Invalid CloudEvent format")


async def _process_grant_alert(notification_data: dict[str, Any]) -> EmailResponse:
    return await send_grant_alert_email(notification_data)  # type: ignore[arg-type]


async def _process_subscription_verification(notification_data: dict[str, Any]) -> EmailResponse:
    return await send_subscription_verification_email(notification_data)  # type: ignore[arg-type]


async def _process_application_email(notification_data: dict[str, Any]) -> EmailResponse:
    application_id = notification_data.get("application_id")
    if not application_id:
        return EmailResponse(status="error", message="Missing application_id in message")

    app_data = await get_application_data(application_id)

    docx_content = markdown_to_docx(app_data["application"]["text"])

    template = jinja_env.get_template("application_ready.html")
    site_url = os.environ.get("SITE_URL", "https://grantflow.ai")

    html_content = template.render(
        application_title=app_data["application"]["title"],
        user_name=app_data["user"]["name"],
        site_url=site_url,
        editor_url=f"{site_url}/projects/{app_data['project']['id']}/applications/{app_data['application']['id']}/editor",
    )

    attachments: list[ResendAttachment] = [
        ResendAttachment(
            filename=f"{app_data['application']['title']}.md",
            content=base64.b64encode(app_data["application"]["text"].encode()).decode(),
        ),
        ResendAttachment(
            filename=f"{app_data['application']['title']}.docx",
            content=base64.b64encode(docx_content).decode(),
        ),
    ]

    subject = f"Your Grant Application is Ready - {app_data['application']['title']}"
    await send_resend_email(app_data["user"]["email"], subject, html_content, attachments)

    logger.info("Email sent for application", application_id=application_id, email=app_data["user"]["email"])

    return EmailResponse(status="success", message="Email sent successfully")


async def send_application_email(cloud_event: CloudEvent) -> EmailResponse:
    try:
        notification_data = _extract_pubsub_data(cloud_event)
        template_type = notification_data.get("template_type")

        if template_type == "grant_alert":
            return await _process_grant_alert(notification_data)
        if template_type == "subscription_verification":
            return await _process_subscription_verification(notification_data)

        return await _process_application_email(notification_data)

    except Exception as e:
        logger.exception("Failed to send application email")
        return EmailResponse(status="error", message=f"Failed to send email: {e!s}")


@functions_framework.cloud_event
def main(cloud_event: CloudEvent) -> None:
    result = asyncio.run(send_application_email(cloud_event))

    if result["status"] == "error":
        logger.error("Email notification error: %s", result["message"])
        raise Exception(result["message"])
    logger.info("Email notification sent: %s", result["message"])
