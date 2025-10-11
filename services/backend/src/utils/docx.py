import re
from io import BytesIO
from typing import TypedDict

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from html_to_markdown import PreprocessingOptions
from html_to_markdown import convert as convert_to_markdown
from markdown import markdown


class TableRow(TypedDict):
    cells: list[str]
    is_header: bool


def markdown_to_docx(markdown_text: str) -> bytes:
    markdown(markdown_text, extensions=["tables", "fenced_code", "nl2br"])

    doc = Document()
    doc.add_heading("Grant Application", 0)

    lines = markdown_text.split("\n")
    current_paragraph = None
    pending_alignment = None
    table_rows: list[TableRow] = []
    in_table = False

    for line in lines:
        stripped_line = line.strip()

        if stripped_line.startswith("<!-- /ALIGNMENT:") and stripped_line.endswith(" -->"):
            current_paragraph = None
            alignment = stripped_line.split("ALIGNMENT:")[1].split(" -->")[0]
            pending_alignment = alignment
            continue

        if "|" in stripped_line and stripped_line.count("|") >= 2:
            if re.match(r"^\|[\s\-:]+\|[\s\-:]*\|", stripped_line):
                if table_rows:
                    last_row = table_rows[-1]
                    if isinstance(last_row, list):
                        table_rows[-1] = {"cells": last_row, "is_header": True}
                    else:
                        table_rows[-1]["is_header"] = True
                in_table = True
                current_paragraph = None
                continue
            cells = [cell.strip() for cell in stripped_line.split("|")[1:-1]]
            table_rows.append({"cells": cells, "is_header": False})
            in_table = True
            current_paragraph = None
            continue
        if in_table and not stripped_line:
            if table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            current_paragraph = None
            continue
        if in_table and not ("|" in stripped_line and stripped_line.count("|") >= 2) and table_rows:
            _add_table_to_doc(doc, table_rows)
            table_rows = []
            in_table = False

        if stripped_line.startswith("# "):
            doc.add_heading(stripped_line[2:], level=1)
            current_paragraph = None
        elif stripped_line.startswith("## "):
            doc.add_heading(stripped_line[3:], level=2)
            current_paragraph = None
        elif stripped_line.startswith("### "):
            doc.add_heading(stripped_line[4:], level=3)
            current_paragraph = None
        elif stripped_line.startswith("#### "):
            doc.add_heading(stripped_line[5:], level=4)
            current_paragraph = None
        elif stripped_line.startswith(("- ", "* ")):
            doc.add_paragraph(stripped_line[2:], style="List Bullet")
            current_paragraph = None
        elif stripped_line.startswith("1. ") or (
            len(stripped_line) > 2 and stripped_line[0].isdigit() and stripped_line[1:3] == ". "
        ):
            doc.add_paragraph(stripped_line[stripped_line.index(". ") + 2 :], style="List Number")
            current_paragraph = None
        elif stripped_line:
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph(stripped_line)
                if pending_alignment:
                    current_paragraph.alignment = _get_paragraph_alignment(pending_alignment)
                    pending_alignment = None
            else:
                current_paragraph.add_run(" " + stripped_line)
        else:
            current_paragraph = None

    if table_rows:
        _add_table_to_doc(doc, table_rows)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_table_to_doc(doc: DocumentObject, table_rows: list[TableRow]) -> None:
    if not table_rows:
        return

    num_rows = len(table_rows)
    num_cols = max(len(row["cells"]) for row in table_rows) if table_rows else 0

    if num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(table_rows):
        cells = row_data["cells"]
        is_header = row_data["is_header"]

        for col_idx in range(num_cols):
            cell_text = cells[col_idx] if col_idx < len(cells) else ""
            cell_obj = table.cell(row_idx, col_idx)
            cell_obj.text = cell_text

            if is_header:
                for paragraph in cell_obj.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _get_paragraph_alignment(alignment: str) -> WD_PARAGRAPH_ALIGNMENT:
    alignment_map = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }

    return alignment_map[alignment]


def html_to_docx(html_content: str) -> bytes:
    markdown_content = convert_to_markdown(
        html_content,
        preprocessing=PreprocessingOptions(enabled=True),
    )
    return markdown_to_docx(markdown_content)
