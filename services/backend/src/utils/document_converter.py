from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from packages.shared_utils.src.exceptions import FileParsingError
from packages.shared_utils.src.logger import get_logger
from weasyprint import CSS, HTML
from weasyprint.text.fonts import FontConfiguration

logger = get_logger(__name__)

DEFAULT_HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.4;
            margin: 0;
            padding: 0.5cm;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 0.8em;
            margin-bottom: 0.3em;
        }}
        h1 {{ font-size: 1.8em; }}
        h2 {{ font-size: 1.4em; }}
        h3 {{ font-size: 1.2em; }}
        p {{
            margin-bottom: 0.6em;
            text-align: left;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 0.5em 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 4px 6px;
            text-align: left;
        }}
        th {{
            background-color: #f8f8f8;
            font-weight: bold;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        .page-break {{
            page-break-before: always;
        }}
        ul, ol {{
            margin: 0.3em 0;
            padding-left: 1.5em;
        }}
        li {{
            margin-bottom: 0.2em;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""

DEFAULT_CSS = """
            @page {
                size: A4;
                margin: 1.5cm;
                @bottom-center {
                    content: "Page " counter(page) " of " counter(pages);
                    font-size: 9pt;
                    color: #999;
                }
            }
            body {
                font-family: Arial, sans-serif;
                line-height: 1.4;
                color: #333;
                margin: 0;
                padding: 0.5cm;
            }
            h1, h2, h3, h4, h5, h6 {
                color: #2c3e50;
                margin-top: 0.8em;
                margin-bottom: 0.3em;
            }
            h1 { font-size: 1.8em; }
            h2 { font-size: 1.4em; }
            h3 { font-size: 1.2em; }
            p {
                margin-bottom: 0.6em;
                text-align: left;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 0.5em 0;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 4px 6px;
                text-align: left;
            }
            th {
                background-color: #f8f8f8;
                font-weight: bold;
            }
            img {
                max-width: 100%;
                height: auto;
            }
            .page-break {
                page-break-before: always;
            }
            ul, ol {
                margin: 0.3em 0;
                padding-left: 1.5em;
            }
            li {
                margin-bottom: 0.2em;
            }
        """


async def convert_html_to_pdf(html_content: str) -> bytes:
    try:
        logger.info(
            "Starting PDF conversion with WeasyPrint",
            html_length=len(html_content),
        )

        html_content = DEFAULT_HTML_TEMPLATE.format(html_content=html_content)
        font_config = FontConfiguration()
        css = CSS(
            string=DEFAULT_CSS,
            font_config=font_config,
        )

        html_doc = HTML(string=html_content)
        pdf_bytes = html_doc.write_pdf(stylesheets=[css], font_config=font_config)

        logger.info(
            "PDF conversion completed successfully",
            pdf_size=len(pdf_bytes),
        )

        return pdf_bytes

    except MemoryError as e:
        logger.error(
            "Memory error during PDF conversion",
            error_type=type(e).__name__,
            error=str(e),
        )
        raise FileParsingError(
            f"Memory error in PDF conversion: {e!s}", context={"error": str(e), "error_type": "memory_error"}
        ) from e


def _parse_html_to_docx(html: str, doc: Document) -> None:
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "li", "table", "div", "span"]):
        text = element.get_text(strip=True)
        if not text:
            continue

        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            heading = doc.add_heading(text, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif element.name == "p":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif element.name in ["ul", "ol"]:
            for li in element.find_all("li", recursive=False):
                li_text = li.get_text(strip=True)
                if li_text:
                    p = doc.add_paragraph(li_text, style="List Bullet")

        elif element.name == "li":
            p = doc.add_paragraph(text, style="List Bullet")

        elif element.name == "table":
            rows = element.find_all("tr")
            if rows:
                max_cols = max(len(row.find_all(["td", "th"])) for row in rows)
                if max_cols > 0:
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = "Table Grid"

                    for i, row in enumerate(rows):
                        cells = row.find_all(["td", "th"])
                        for j, cell in enumerate(cells):
                            if j < len(table.rows[i].cells):
                                cell_text = cell.get_text(strip=True)
                                table.rows[i].cells[j].text = cell_text

        elif element.name in ["div", "span"] and element.parent.name not in ["p", "li", "td", "th"]:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT


async def convert_html_to_docx(html_content: str) -> bytes:
    try:
        logger.info(
            "Starting DOCX conversion with python-docx",
            html_length=len(html_content),
        )

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        _parse_html_to_docx(html_content, doc)

        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        logger.info(
            "DOCX conversion completed successfully",
            docx_size=len(docx_bytes.getvalue()),
        )

        return docx_bytes.getvalue()

    except MemoryError as e:
        logger.error(
            "Memory error during DOCX conversion",
            error_type=type(e).__name__,
            error=str(e),
        )
        raise FileParsingError(
            f"Memory error in DOCX conversion: {e!s}", context={"error": str(e), "error_type": "memory_error"}
        ) from e
