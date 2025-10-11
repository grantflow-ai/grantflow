from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from services.backend.src.utils.docx import html_to_docx, markdown_to_docx


@pytest.fixture
def sample_markdown() -> str:
    return """# Main Title

This is a paragraph with some text.

## Section Heading

Here's another paragraph with more content.

### Subsection

#### Subsection Alt

- Bullet point 1
- Bullet point 2
* Another bullet point

1. Numbered item 1
2. Numbered item 2
3. Numbered item 3

This is a final paragraph."""


def test_markdown_to_docx_creates_document(sample_markdown: str) -> None:
    result = markdown_to_docx(sample_markdown)

    assert isinstance(result, bytes)
    assert len(result) > 0

    buffer = BytesIO(result)
    doc = Document(buffer)

    assert doc is not None
    assert len(doc.paragraphs) > 0


def test_markdown_to_docx_preserves_headings(sample_markdown: str) -> None:
    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(sample_markdown)

        assert mock_doc.add_heading.call_count >= 4
        mock_doc.add_heading.assert_any_call("Grant Application", 0)
        mock_doc.add_heading.assert_any_call("Main Title", level=1)
        mock_doc.add_heading.assert_any_call("Section Heading", level=2)
        mock_doc.add_heading.assert_any_call("Subsection", level=3)
        mock_doc.add_heading.assert_any_call("Subsection Alt", level=4)


def test_markdown_to_docx_handles_lists() -> None:
    markdown_text = """# Title

- Bullet item 1
- Bullet item 2
* Star bullet item

1. Number item 1
2. Number item 2"""

    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(markdown_text)

        bullet_calls = [
            call
            for call in mock_doc.add_paragraph.call_args_list
            if call.args and len(call.args) > 0 and call.kwargs.get("style") == "List Bullet"
        ]
        assert len(bullet_calls) == 3

        number_calls = [
            call
            for call in mock_doc.add_paragraph.call_args_list
            if call.args and len(call.args) > 0 and call.kwargs.get("style") == "List Number"
        ]
        assert len(number_calls) == 2


def test_markdown_to_docx_handles_paragraphs() -> None:
    markdown_text = """# Title

First paragraph text.

Second paragraph text that is longer
and continues on multiple lines.

Third paragraph."""

    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_paragraph = MagicMock()
        mock_doc.add_paragraph.return_value = mock_paragraph
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(markdown_text)

        regular_paragraph_calls = [
            call
            for call in mock_doc.add_paragraph.call_args_list
            if call.args and len(call.args) > 0 and "style" not in call.kwargs
        ]
        assert len(regular_paragraph_calls) >= 2

        assert mock_paragraph.add_run.called


def test_markdown_to_docx_handles_empty_lines() -> None:
    markdown_text = """# Title


Paragraph with empty lines above.


Another paragraph."""

    result = markdown_to_docx(markdown_text)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_markdown_to_docx_handles_edge_cases() -> None:
    edge_cases = [
        "",
        "# ",
        "## ",
        "- ",
        "1. ",
        "\n\n\n",
        "Simple text without formatting",
    ]

    for markdown_text in edge_cases:
        result = markdown_to_docx(markdown_text)
        assert isinstance(result, bytes)
        assert len(result) > 0


def test_markdown_to_docx_single_digit_numbered_lists() -> None:
    markdown_text = """1. First item
2. Second item
9. Ninth item"""

    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(markdown_text)

        number_calls = [
            call
            for call in mock_doc.add_paragraph.call_args_list
            if call.args and len(call.args) > 0 and call.kwargs.get("style") == "List Number"
        ]
        assert len(number_calls) == 3

        number_texts = [call.args[0] for call in number_calls]
        assert "First item" in number_texts
        assert "Second item" in number_texts
        assert "Ninth item" in number_texts


@pytest.mark.parametrize(
    "text,expected_alignment",
    [
        (
            """<!-- /ALIGNMENT:left -->
This text is aligned left""",
            WD_PARAGRAPH_ALIGNMENT.LEFT,
        ),
        (
            """<!-- /ALIGNMENT:center -->
This text is aligned left""",
            WD_PARAGRAPH_ALIGNMENT.CENTER,
        ),
        (
            """<!-- /ALIGNMENT:right -->
This text is aligned left""",
            WD_PARAGRAPH_ALIGNMENT.RIGHT,
        ),
        (
            """<!-- /ALIGNMENT:justify -->
This text is aligned left""",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        ),
    ],
)
def test_markdown_to_docx_handles_text_alignment(text: str, expected_alignment: WD_PARAGRAPH_ALIGNMENT) -> None:
    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_paragraph = MagicMock()
        mock_doc.add_paragraph.return_value = mock_paragraph
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(text)
        assert mock_paragraph.alignment == expected_alignment


def test_markdown_to_docx_handles_tables() -> None:
    table_markdown = """# Title

| Name | Age | City |
|------|-----|------|
| John | 25  | NYC  |
| Jane | 30  | LA   |

Some text after table."""

    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_doc.add_table.return_value = mock_table
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(table_markdown)

        mock_doc.add_table.assert_called_once_with(rows=3, cols=3)

        assert mock_table.style == "Table Grid"

        assert mock_table.cell.call_count == 9

        mock_table.cell.assert_any_call(0, 0)
        mock_table.cell.assert_any_call(0, 1)
        mock_table.cell.assert_any_call(0, 2)
        mock_table.cell.assert_any_call(1, 0)
        mock_table.cell.assert_any_call(1, 1)
        mock_table.cell.assert_any_call(1, 2)
        mock_table.cell.assert_any_call(2, 0)
        mock_table.cell.assert_any_call(2, 1)
        mock_table.cell.assert_any_call(2, 2)


def test_markdown_to_docx_handles_table_without_header() -> None:
    table_markdown = """# Title

| John | 25  | NYC  |
| Jane | 30  | LA   |

Some text after table."""

    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_doc.add_table.return_value = mock_table
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(table_markdown)

        mock_doc.add_table.assert_called_once_with(rows=2, cols=3)

        assert mock_table.style == "Table Grid"


def test_markdown_to_docx_handles_uneven_table_rows() -> None:
    table_markdown = """# Title

| Name | Age | City |
|------|-----|------|
| John | 25  |
| Jane | 30  | LA   | Boston |

Some text after table."""

    with patch("services.backend.src.utils.docx.Document") as mock_doc_class:
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_doc.add_table.return_value = mock_table
        mock_doc_class.return_value = mock_doc

        BytesIO()
        mock_doc.save = lambda buffer: buffer.write(b"test")

        markdown_to_docx(table_markdown)

        mock_doc.add_table.assert_called_once_with(rows=3, cols=4)

        assert mock_table.style == "Table Grid"


def test_html_to_docx_creates_document() -> None:
    html_content = """
    <h1>Main Title</h1>
    <p>This is a paragraph with some text.</p>
    <h2>Section Heading</h2>
    <p>Here's another paragraph with more content.</p>
    <ul>
        <li>Bullet point 1</li>
        <li>Bullet point 2</li>
    </ul>
    <ol>
        <li>Numbered item 1</li>
        <li>Numbered item 2</li>
    </ol>
    """

    result = html_to_docx(html_content)

    assert isinstance(result, bytes)
    assert len(result) > 0

    buffer = BytesIO(result)
    doc = Document(buffer)

    assert doc is not None
    assert len(doc.paragraphs) > 0
